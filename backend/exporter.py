import json
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO


def export_to_docx(session_filename: str, answers: list) -> bytes:
    doc = Document()

    doc.add_heading(f"Answered Questionnaire: {session_filename}", level=1)

    total = len(answers)
    answered = sum(
        1 for a in answers
        if a.answer and "not found" not in (a.answer or "").lower()
    )
    not_found = total - answered

    summary = doc.add_paragraph()
    summary.add_run(
        f"Coverage: {answered}/{total} answered  |  {not_found} not found in references"
    ).font.size = Pt(10)
    doc.add_paragraph("")

    for a in answers:
        # ── Question ──────────────────────────────────────────────────────────
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{a.index + 1}. {a.question}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # ── Answer ────────────────────────────────────────────────────────────
        answer_text = a.answer or "—"
        is_not_found = "not found" in answer_text.lower()

        a_para = doc.add_paragraph()
        a_run = a_para.add_run(f"Answer: {answer_text}")
        a_run.font.size = Pt(10)
        if is_not_found:
            a_run.font.color.rgb = RGBColor(0xE5, 0x39, 0x35)

        if a.is_edited:
            edited_run = a_para.add_run("  (edited)")
            edited_run.font.size = Pt(9)
            edited_run.font.color.rgb = RGBColor(0x81, 0x8C, 0xF8)

        # ── Confidence ────────────────────────────────────────────────────────
        if a.confidence and a.confidence > 0 and not is_not_found:
            conf_para = doc.add_paragraph()
            conf_run = conf_para.add_run(
                f"Confidence: {round(a.confidence * 100)}%"
            )
            conf_run.font.size = Pt(9)
            conf_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

        # ── Evidence snippet ──────────────────────────────────────────────────
        try:
            snippets = json.loads(a.evidence_snippets or "[]")
        except Exception:
            snippets = []

        if snippets and not is_not_found:
            ev_para = doc.add_paragraph()
            ev_run = ev_para.add_run(f'Evidence: "{snippets[0]}"')
            ev_run.font.size = Pt(9)
            ev_run.italic = True
            ev_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

        # ── Citations ─────────────────────────────────────────────────────────
        try:
            citations = json.loads(a.citations or "[]")
        except Exception:
            citations = []

        if citations:
            cit_para = doc.add_paragraph()
            cit_run = cit_para.add_run("Citations: " + "  |  ".join(citations))
            cit_run.font.size = Pt(9)
            cit_run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

        doc.add_paragraph("")  # spacer between answers

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
